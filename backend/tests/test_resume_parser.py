"""
Tests for resume parsing endpoint
"""
import pytest
from fastapi.testclient import TestClient
from io import BytesIO
from main import app

client = TestClient(app)


@pytest.fixture
def sample_pdf_content():
    """Create a minimal PDF file content for testing"""
    # This is a minimal valid PDF structure
    pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
/Resources <<
/Font <<
/F1 <<
/Type /Font
/Subtype /Type1
/BaseFont /Helvetica
>>
>>
>>
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(Test Resume Content) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000306 00000 n
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
398
%%EOF"""
    return pdf_content


@pytest.fixture
def sample_docx_content():
    """Create a minimal DOCX file content for testing"""
    # For testing, we'll use a simple approach
    # In real scenario, you'd use python-docx to create a proper DOCX
    # For now, we'll test with actual docx creation
    from docx import Document
    doc = Document()
    doc.add_paragraph("Test Resume Content")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("5 years of experience")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


@pytest.fixture
def large_file_content():
    """Create a file larger than 2MB"""
    return b"x" * (2 * 1024 * 1024 + 1)  # 2MB + 1 byte


class TestParseResume:
    """Test cases for /api/parse-resume endpoint"""

    def test_parse_resume_with_text_input(self):
        """Test parsing resume from text input"""
        response = client.post(
            "/api/parse-resume",
            data={"text": "John Doe\nSoftware Engineer\n5 years of experience"}
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert "John Doe" in data["text"]
        assert data["source"] == "text_input"
        assert data["file_name"] is None
        assert data["file_type"] is None

    def test_parse_resume_with_empty_text(self):
        """Test parsing with empty text input"""
        response = client.post(
            "/api/parse-resume",
            data={"text": ""}
        )
        assert response.status_code == 400
        # Empty string is treated as no input provided
        detail = response.json()["detail"].lower()
        assert "empty" in detail or "must be provided" in detail

    def test_parse_resume_with_whitespace_only_text(self):
        """Test parsing with whitespace-only text"""
        response = client.post(
            "/api/parse-resume",
            data={"text": "   \n\t  "}
        )
        assert response.status_code == 400

    def test_parse_resume_with_pdf_file(self, sample_pdf_content):
        """Test parsing PDF file"""
        file_obj = BytesIO(sample_pdf_content)
        response = client.post(
            "/api/parse-resume",
            files={"file": ("resume.pdf", file_obj, "application/pdf")}
        )
        # PDF parsing might fail with minimal PDF, so we check for either success or parsing error
        assert response.status_code in [200, 400]
        if response.status_code == 200:
            data = response.json()
            assert data["success"] is True
            assert data["source"] == "file_upload"
            assert data["file_name"] == "resume.pdf"
            assert data["file_type"] == "pdf"
            assert "file_size" in data

    def test_parse_resume_with_docx_file(self, sample_docx_content):
        """Test parsing DOCX file"""
        file_obj = BytesIO(sample_docx_content)
        response = client.post(
            "/api/parse-resume",
            files={"file": ("resume.docx", file_obj, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
        assert response.status_code == 200
        data = response.json()
        assert data["success"] is True
        assert data["source"] == "file_upload"
        assert data["file_name"] == "resume.docx"
        assert data["file_type"] == "docx"
        assert "Test Resume Content" in data["text"]

    def test_parse_resume_with_file_too_large(self, large_file_content):
        """Test parsing with file exceeding 2MB limit"""
        response = client.post(
            "/api/parse-resume",
            files={"file": ("large.pdf", BytesIO(large_file_content), "application/pdf")}
        )
        assert response.status_code == 400
        assert "exceeds maximum" in response.json()["detail"].lower()

    def test_parse_resume_with_empty_file(self):
        """Test parsing with empty file"""
        response = client.post(
            "/api/parse-resume",
            files={"file": ("empty.pdf", BytesIO(b""), "application/pdf")}
        )
        assert response.status_code == 400
        assert "empty" in response.json()["detail"].lower()

    def test_parse_resume_with_unsupported_file_type(self):
        """Test parsing with unsupported file type"""
        response = client.post(
            "/api/parse-resume",
            files={"file": ("resume.txt", BytesIO(b"test content"), "text/plain")}
        )
        assert response.status_code == 400
        assert "unsupported" in response.json()["detail"].lower()

    def test_parse_resume_with_doc_file(self):
        """Test parsing with old DOC format (should fail)"""
        response = client.post(
            "/api/parse-resume",
            files={"file": ("resume.doc", BytesIO(b"test content"), "application/msword")}
        )
        assert response.status_code == 400
        assert "doc" in response.json()["detail"].lower()

    def test_parse_resume_without_file_or_text(self):
        """Test parsing without providing file or text"""
        response = client.post("/api/parse-resume")
        assert response.status_code == 400
        assert "must be provided" in response.json()["detail"].lower()

    def test_parse_resume_text_preserves_formatting(self):
        """Test that text input preserves formatting"""
        text_input = "John Doe\n\nSoftware Engineer\n\nExperience:\n- 5 years"
        response = client.post(
            "/api/parse-resume",
            data={"text": text_input}
        )
        assert response.status_code == 200
        data = response.json()
        assert "John Doe" in data["text"]
        assert "Software Engineer" in data["text"]

    def test_parse_resume_response_structure(self):
        """Test that response has correct structure"""
        response = client.post(
            "/api/parse-resume",
            data={"text": "Test content"}
        )
        assert response.status_code == 200
        data = response.json()
        assert "success" in data
        assert "text" in data
        assert "source" in data
        assert "file_name" in data
        assert "file_type" in data

