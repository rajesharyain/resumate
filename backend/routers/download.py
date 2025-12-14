"""
Resume download router - PDF and DOCX generation
"""
from fastapi import APIRouter, HTTPException, Body
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from typing import Dict, Any, List, Optional
import io
import html

# Try to import weasyprint, but make it optional for Windows compatibility
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    WEASYPRINT_ERROR = str(e)

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# StructuredResume type is defined in convert.py but we use Dict here for flexibility

router = APIRouter(prefix="/api", tags=["download"])


def escape_html(text: str) -> str:
    """Escape HTML special characters"""
    if not text:
        return ""
    return html.escape(str(text))


class DownloadRequest(BaseModel):
    resume: Dict[str, Any]  # StructuredResume as dict
    standard: str


def generate_resume_html(resume_data: Dict[str, Any], standard: str) -> str:
    """Generate HTML for resume based on standard"""
    
    # Get template based on standard
    if standard == "us_ats":
        return _generate_us_ats_html(resume_data)
    elif standard == "europass":
        return _generate_europass_html(resume_data)
    elif standard == "indian_corporate":
        return _generate_indian_corporate_html(resume_data)
    elif standard == "uk_professional":
        return _generate_uk_professional_html(resume_data)
    else:
        return _generate_us_ats_html(resume_data)


def _generate_us_ats_html(resume: Dict[str, Any]) -> str:
    """Generate US ATS HTML"""
    personal_info = resume.get("personal_info", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 8.5in;
                margin: 0 auto;
                padding: 0.5in;
                color: #333;
                line-height: 1.6;
            }}
            header {{
                text-align: center;
                border-bottom: 2px solid #333;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            h1 {{
                font-size: 28px;
                margin: 10px 0;
                color: #000;
            }}
            .contact-info {{
                font-size: 12px;
                color: #666;
                margin-top: 10px;
            }}
            .contact-info span {{
                margin: 0 10px;
            }}
            section {{
                margin-bottom: 25px;
            }}
            h2 {{
                font-size: 18px;
                border-bottom: 1px solid #333;
                padding-bottom: 5px;
                margin-bottom: 15px;
                color: #000;
            }}
            .experience-item, .education-item {{
                margin-bottom: 20px;
            }}
            .experience-header, .education-header {{
                display: flex;
                justify-content: space-between;
                margin-bottom: 5px;
            }}
            .job-title {{
                font-weight: bold;
                font-size: 16px;
            }}
            .company {{
                font-weight: 600;
                color: #555;
            }}
            .date-location {{
                font-size: 12px;
                color: #666;
                text-align: right;
            }}
            ul {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            li {{
                margin: 5px 0;
            }}
            .skills {{
                display: flex;
                flex-wrap: wrap;
                gap: 8px;
            }}
            .skill-tag {{
                background-color: #f0f0f0;
                padding: 4px 12px;
                border-radius: 4px;
                font-size: 12px;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
            }}
        </style>
    </head>
    <body>
        <header>
            <h1>{personal_info.get('full_name', 'Your Name')}</h1>
            <div class="contact-info">
                {f"<span>{personal_info.get('email', '')}</span>" if personal_info.get('email') else ""}
                {f"<span>{personal_info.get('phone', '')}</span>" if personal_info.get('phone') else ""}
                {f"<span>{personal_info.get('location', '')}</span>" if personal_info.get('location') else ""}
            </div>
        </header>
    """
    
    if summary:
        html += f"""
        <section>
            <h2>Professional Summary</h2>
            <p>{summary}</p>
        </section>
        """
    
    if experience:
        html_content += """
        <section>
            <h2>Work Experience</h2>
        """
        for exp in experience:
            title = html.escape(exp.get('title', 'Job Title'))
            company = html.escape(exp.get('company', 'Company'))
            exp_location = html.escape(exp.get('location', 'Location'))
            start_date = html.escape(exp.get('start_date', 'Start'))
            end_date = html.escape(exp.get('end_date', 'End'))
            
            html_content += f"""
            <div class="experience-item">
                <div class="experience-header">
                    <div>
                        <div class="job-title">{title}</div>
                        <div class="company">{company}</div>
                    </div>
                    <div class="date-location">
                        <div>{exp_location}</div>
                        <div>{start_date} - {end_date}</div>
                    </div>
                </div>
            """
            if exp.get('achievements'):
                html_content += "<ul>"
                for achievement in exp['achievements']:
                    achievement_escaped = html.escape(achievement)
                    html_content += f"<li>{achievement_escaped}</li>"
                html_content += "</ul>"
            html_content += "</div>"
        html_content += "</section>"
    
    if education:
        html_content += """
        <section>
            <h2>Education</h2>
        """
        for edu in education:
            degree = html.escape(edu.get('degree', 'Degree'))
            institution = html.escape(edu.get('institution', 'Institution'))
            edu_location = html.escape(edu.get('location', 'Location'))
            graduation_date = html.escape(edu.get('graduation_date', 'Year'))
            
            html_content += f"""
            <div class="education-item">
                <div class="education-header">
                    <div>
                        <div class="job-title">{degree}</div>
                        <div class="company">{institution}</div>
                        <div style="font-size: 12px; color: #666;">{edu_location}</div>
                    </div>
                    <div class="date-location">{graduation_date}</div>
                </div>
            </div>
            """
        html_content += "</section>"
    
    if skills and isinstance(skills, list) and len(skills) > 0:
        html_content += """
        <section>
            <h2>Skills</h2>
            <div class="skills">
        """
        for skill in skills:
            if isinstance(skill, str):
                skill_escaped = html.escape(skill)
                html_content += f'<span class="skill-tag">{skill_escaped}</span>'
        html_content += """
            </div>
        </section>
        """
    
    html_content += """
    </body>
    </html>
    """
    return html_content


def _generate_europass_html(resume: Dict[str, Any]) -> str:
    """Generate EUROPASS HTML"""
    personal_info = resume.get("personal_info", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 8.5in;
                margin: 0 auto;
                padding: 0.5in;
                color: #333;
                line-height: 1.6;
            }}
            header {{
                border-bottom: 3px solid #0066CC;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            h1 {{
                font-size: 28px;
                margin: 10px 0;
                color: #0066CC;
            }}
            .contact-info {{
                font-size: 12px;
                color: #666;
                margin-top: 10px;
            }}
            h2 {{
                font-size: 18px;
                color: #0066CC;
                border-bottom: 2px solid #0066CC;
                padding-bottom: 5px;
                margin-bottom: 15px;
            }}
            .experience-item, .education-item {{
                margin-bottom: 20px;
            }}
            .experience-header, .education-header {{
                display: flex;
                justify-content: space-between;
                margin-bottom: 5px;
            }}
            .job-title {{
                font-weight: bold;
                font-size: 16px;
            }}
            ul {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            .skills {{
                display: flex;
                flex-wrap: wrap;
                gap: 8px;
            }}
            .skill-tag {{
                background-color: #E6F2FF;
                padding: 4px 12px;
                border-radius: 4px;
                font-size: 12px;
                border: 1px solid #0066CC;
            }}
        </style>
    </head>
    <body>
        <header>
            <h1>{personal_info.get('full_name', 'Your Name')}</h1>
            <div class="contact-info">
                {f"Email: {personal_info.get('email', '')}<br>" if personal_info.get('email') else ""}
                {f"Phone: {personal_info.get('phone', '')}<br>" if personal_info.get('phone') else ""}
                {f"Address: {personal_info.get('location', '')}" if personal_info.get('location') else ""}
            </div>
        </header>
    """
    
    if summary:
        html += f"""
        <section>
            <h2>Personal Statement</h2>
            <p>{summary}</p>
        </section>
        """
    
    if experience:
        html += """
        <section>
            <h2>Work Experience</h2>
        """
        for exp in experience:
            html += f"""
            <div class="experience-item">
                <div class="experience-header">
                    <div>
                        <div class="job-title">{exp.get('title', 'Job Title')}</div>
                        <div>{exp.get('company', 'Company')}</div>
                        {f"<div style='font-size: 12px; color: #666;'>{exp.get('description', '')}</div>" if exp.get('description') else ""}
                    </div>
                    <div style="font-size: 12px; color: #666; text-align: right;">
                        <div>{exp.get('location', 'Location')}</div>
                        <div>{exp.get('start_date', 'Start')} - {exp.get('end_date', 'End')}</div>
                    </div>
                </div>
            """
            if exp.get('achievements'):
                html += "<ul>"
                for achievement in exp['achievements']:
                    html += f"<li>{achievement}</li>"
                html += "</ul>"
            html += "</div>"
        html += "</section>"
    
    if education:
        html += """
        <section>
            <h2>Education and Training</h2>
        """
        for edu in education:
            html += f"""
            <div class="education-item">
                <div class="education-header">
                    <div>
                        <div class="job-title">{edu.get('degree', 'Degree')}</div>
                        {f"<div style='font-size: 12px;'>{edu.get('field_of_study', '')}</div>" if edu.get('field_of_study') else ""}
                        <div>{edu.get('institution', 'Institution')}</div>
                        <div style="font-size: 12px; color: #666;">{edu.get('location', 'Location')}</div>
                    </div>
                    <div style="font-size: 12px; color: #666; text-align: right;">
                        <div>{edu.get('graduation_date', 'Year')}</div>
                        {f"<div>{edu.get('grade', '')}</div>" if edu.get('grade') else ""}
                    </div>
                </div>
            </div>
            """
        html += "</section>"
    
    if skills and isinstance(skills, list) and len(skills) > 0:
        html += """
        <section>
            <h2>Skills and Competences</h2>
            <div class="skills">
        """
        for skill in skills:
            if isinstance(skill, str):
                html += f'<span class="skill-tag">{skill}</span>'
        html += """
            </div>
        </section>
        """
    
    html += """
    </body>
    </html>
    """
    return html


def _generate_indian_corporate_html(resume: Dict[str, Any]) -> str:
    """Generate Indian Corporate HTML"""
    personal_info = resume.get("personal_info", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 8.5in;
                margin: 0 auto;
                padding: 0.5in;
                color: #333;
                line-height: 1.6;
            }}
            header {{
                border-bottom: 3px solid #4F46E5;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            h1 {{
                font-size: 28px;
                margin: 10px 0;
                color: #4F46E5;
            }}
            .contact-info {{
                font-size: 12px;
                color: #666;
                margin-top: 10px;
            }}
            h2 {{
                font-size: 18px;
                color: #4F46E5;
                border-bottom: 2px solid #4F46E5;
                padding-bottom: 5px;
                margin-bottom: 15px;
            }}
            .experience-item, .education-item {{
                margin-bottom: 20px;
            }}
            .job-title {{
                font-weight: bold;
                font-size: 16px;
            }}
            ul {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            .skills {{
                display: flex;
                flex-wrap: wrap;
                gap: 8px;
            }}
            .skill-tag {{
                background-color: #EEF2FF;
                padding: 4px 12px;
                border-radius: 4px;
                font-size: 12px;
                border: 1px solid #4F46E5;
            }}
        </style>
    </head>
    <body>
        <header>
            <h1>{personal_info.get('full_name', 'Your Name')}</h1>
            <div class="contact-info">
                {f"Email: {personal_info.get('email', '')}<br>" if personal_info.get('email') else ""}
                {f"Mobile: {personal_info.get('phone', '')}<br>" if personal_info.get('phone') else ""}
                {f"Location: {personal_info.get('location', '')}<br>" if personal_info.get('location') else ""}
                {f"Current CTC: {personal_info.get('current_ctc', '')}<br>" if personal_info.get('current_ctc') else ""}
                {f"Expected CTC: {personal_info.get('expected_ctc', '')}<br>" if personal_info.get('expected_ctc') else ""}
                {f"Notice Period: {personal_info.get('notice_period', '')}" if personal_info.get('notice_period') else ""}
            </div>
        </header>
    """
    
    if summary:
        html += f"""
        <section>
            <h2>Professional Summary</h2>
            <p>{summary}</p>
        </section>
        """
    
    if experience:
        html += """
        <section>
            <h2>Professional Experience</h2>
        """
        for exp in experience:
            html += f"""
            <div class="experience-item">
                <div style="display: flex; justify-content: space-between; margin-bottom: 5px;">
                    <div>
                        <div class="job-title">{exp.get('title', 'Job Title')}</div>
                        <div>{exp.get('company', 'Company')}</div>
                        <div style="font-size: 12px; color: #666;">{exp.get('location', 'Location')}</div>
                    </div>
                    <div style="font-size: 12px; color: #666; text-align: right;">
                        {exp.get('start_date', 'Start')} - {exp.get('end_date', 'End')}
                    </div>
                </div>
            """
            if exp.get('achievements'):
                html += "<ul>"
                for achievement in exp['achievements']:
                    html += f"<li>{achievement}</li>"
                html += "</ul>"
            html += "</div>"
        html += "</section>"
    
    if education:
        html += """
        <section>
            <h2>Education</h2>
        """
        for edu in education:
            html += f"""
            <div class="education-item">
                <div style="display: flex; justify-content: space-between;">
                    <div>
                        <div class="job-title">{edu.get('degree', 'Degree')}</div>
                        <div>{edu.get('institution', 'Institution')}</div>
                        {f"<div style='font-size: 12px;'>{edu.get('university', '')}</div>" if edu.get('university') else ""}
                        <div style="font-size: 12px; color: #666;">{edu.get('location', 'Location')}</div>
                    </div>
                    <div style="font-size: 12px; color: #666; text-align: right;">
                        <div>{edu.get('graduation_date', 'Year')}</div>
                        {f"<div>{edu.get('percentage', '')}</div>" if edu.get('percentage') else ""}
                    </div>
                </div>
            </div>
            """
        html += "</section>"
    
    if skills and isinstance(skills, list) and len(skills) > 0:
        html += """
        <section>
            <h2>Technical Skills</h2>
            <div class="skills">
        """
        for skill in skills:
            if isinstance(skill, str):
                html += f'<span class="skill-tag">{skill}</span>'
        html += """
            </div>
        </section>
        """
    
    html += """
    </body>
    </html>
    """
    return html


def _generate_uk_professional_html(resume: Dict[str, Any]) -> str:
    """Generate UK Professional HTML"""
    personal_info = resume.get("personal_info", {})
    summary = resume.get("summary", "")
    experience = resume.get("experience", [])
    education = resume.get("education", [])
    skills = resume.get("skills", [])
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                max-width: 8.5in;
                margin: 0 auto;
                padding: 0.5in;
                color: #333;
                line-height: 1.6;
            }}
            header {{
                border-bottom: 3px solid #475569;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            h1 {{
                font-size: 28px;
                margin: 10px 0;
                color: #1E293B;
            }}
            .contact-info {{
                font-size: 12px;
                color: #666;
                margin-top: 10px;
            }}
            h2 {{
                font-size: 18px;
                color: #475569;
                border-bottom: 2px solid #475569;
                padding-bottom: 5px;
                margin-bottom: 15px;
            }}
            .experience-item, .education-item {{
                margin-bottom: 20px;
            }}
            .job-title {{
                font-weight: bold;
                font-size: 16px;
            }}
            ul {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            .skills {{
                display: flex;
                flex-wrap: wrap;
                gap: 8px;
            }}
            .skill-tag {{
                background-color: #F1F5F9;
                padding: 4px 12px;
                border-radius: 4px;
                font-size: 12px;
                border: 1px solid #475569;
            }}
        </style>
    </head>
    <body>
        <header>
            <h1>{personal_info.get('full_name', 'Your Name')}</h1>
            <div class="contact-info">
                {f"Email: {personal_info.get('email', '')}<br>" if personal_info.get('email') else ""}
                {f"Telephone: {personal_info.get('phone', '')}<br>" if personal_info.get('phone') else ""}
                {f"Location: {personal_info.get('location', '')}" if personal_info.get('location') else ""}
            </div>
        </header>
    """
    
    if summary:
        html += f"""
        <section>
            <h2>Professional Profile</h2>
            <p>{summary}</p>
        </section>
        """
    
    if experience:
        html += """
        <section>
            <h2>Professional Experience</h2>
        """
        for exp in experience:
            html += f"""
            <div class="experience-item">
                <div style="display: flex; justify-content: space-between; margin-bottom: 5px;">
                    <div>
                        <div class="job-title">{exp.get('title', 'Job Title')}</div>
                        <div>{exp.get('company', 'Company')}</div>
                    </div>
                    <div style="font-size: 12px; color: #666; text-align: right;">
                        <div>{exp.get('location', 'Location')}</div>
                        <div>{exp.get('start_date', 'Start')} - {exp.get('end_date', 'End')}</div>
                    </div>
                </div>
            """
            if exp.get('achievements'):
                html += "<ul>"
                for achievement in exp['achievements']:
                    html += f"<li>{achievement}</li>"
                html += "</ul>"
            html += "</div>"
        html += "</section>"
    
    if education:
        html += """
        <section>
            <h2>Education and Qualifications</h2>
        """
        for edu in education:
            html += f"""
            <div class="education-item">
                <div style="display: flex; justify-content: space-between;">
                    <div>
                        <div class="job-title">{edu.get('degree', 'Degree')}</div>
                        <div>{edu.get('institution', 'Institution')}</div>
                        <div style="font-size: 12px; color: #666;">{edu.get('location', 'Location')}</div>
                    </div>
                    <div style="font-size: 12px; color: #666; text-align: right;">
                        <div>{edu.get('graduation_date', 'Year')}</div>
                        {f"<div>{edu.get('grade', '')}</div>" if edu.get('grade') else ""}
                    </div>
                </div>
            </div>
            """
        html += "</section>"
    
    if skills and isinstance(skills, list) and len(skills) > 0:
        html += """
        <section>
            <h2>Key Skills</h2>
            <div class="skills">
        """
        for skill in skills:
            if isinstance(skill, str):
                html += f'<span class="skill-tag">{skill}</span>'
        html += """
            </div>
        </section>
        """
    
    html += """
    </body>
    </html>
    """
    return html


def generate_resume_docx(resume_data: Dict[str, Any], standard: str) -> Document:
    """Generate DOCX document from resume data"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    personal_info = resume_data.get("personal_info", {})
    summary = resume_data.get("summary", "")
    experience = resume_data.get("experience", [])
    education = resume_data.get("education", [])
    skills = resume_data.get("skills", [])
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(personal_info.get('full_name', 'Your Name'))
    run.bold = True
    run.font.size = Pt(20)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_info = []
    if personal_info.get('email'):
        contact_info.append(personal_info['email'])
    if personal_info.get('phone'):
        contact_info.append(personal_info['phone'])
    if personal_info.get('location'):
        contact_info.append(personal_info['location'])
    contact.add_run(' | '.join(contact_info))
    contact.runs[0].font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Summary
    if summary:
        doc.add_paragraph('Professional Summary').runs[0].bold = True
        doc.add_paragraph(summary)
        doc.add_paragraph()  # Spacing
    
    # Experience
    if experience:
        doc.add_paragraph('Work Experience').runs[0].bold = True
        for exp in experience:
            exp_para = doc.add_paragraph()
            exp_para.add_run(exp.get('title', 'Job Title')).bold = True
            exp_para.add_run(f" | {exp.get('company', 'Company')}")
            exp_para.add_run(f" | {exp.get('start_date', 'Start')} - {exp.get('end_date', 'End')}")
            
            if exp.get('achievements'):
                for achievement in exp['achievements']:
                    doc.add_paragraph(achievement, style='List Bullet')
            doc.add_paragraph()  # Spacing
    
    # Education
    if education:
        doc.add_paragraph('Education').runs[0].bold = True
        for edu in education:
            edu_para = doc.add_paragraph()
            edu_para.add_run(edu.get('degree', 'Degree')).bold = True
            edu_para.add_run(f" | {edu.get('institution', 'Institution')}")
            edu_para.add_run(f" | {edu.get('graduation_date', 'Year')}")
            doc.add_paragraph()  # Spacing
    
    # Skills
    if skills and isinstance(skills, list) and len(skills) > 0:
        doc.add_paragraph('Skills').runs[0].bold = True
        skills_text = ', '.join([s for s in skills if isinstance(s, str)])
        doc.add_paragraph(skills_text)
    
    return doc


@router.post("/download/pdf")
async def download_pdf(request: DownloadRequest):
    """
    Generate and download resume as PDF.
    
    Accepts structured resume data and returns PDF file.
    """
    # Check if WeasyPrint is available
    if not WEASYPRINT_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail=f"PDF generation is not available on this system. WeasyPrint requires system libraries that are not installed. Error: {WEASYPRINT_ERROR if 'WEASYPRINT_ERROR' in globals() else 'WeasyPrint not available'}. Please use DOCX download instead or install WeasyPrint dependencies."
        )
    
    try:
        # Validate input
        if not request.resume:
            raise HTTPException(
                status_code=400,
                detail="Resume data is required"
            )
        
        # Generate HTML
        html_content = generate_resume_html(request.resume, request.standard)
        
        # Convert HTML to PDF
        try:
            pdf_bytes = HTML(string=html_content).write_pdf()
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to generate PDF: {str(e)}"
            )
        
        # Return PDF file
        filename = f"resume_{request.standard}_{request.resume.get('personal_info', {}).get('full_name', 'resume').replace(' ', '_')}.pdf"
        
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred while generating PDF: {str(e)}"
        )


@router.post("/download/docx")
async def download_docx(request: DownloadRequest):
    """
    Generate and download resume as DOCX.
    
    Accepts structured resume data and returns DOCX file.
    """
    try:
        # Validate input
        if not request.resume:
            raise HTTPException(
                status_code=400,
                detail="Resume data is required"
            )
        
        # Generate DOCX
        try:
            doc = generate_resume_docx(request.resume, request.standard)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to generate DOCX: {str(e)}"
            )
        
        # Save to bytes
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        # Return DOCX file
        filename = f"resume_{request.standard}_{request.resume.get('personal_info', {}).get('full_name', 'resume').replace(' ', '_')}.docx"
        
        return Response(
            content=docx_bytes.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred while generating DOCX: {str(e)}"
        )

